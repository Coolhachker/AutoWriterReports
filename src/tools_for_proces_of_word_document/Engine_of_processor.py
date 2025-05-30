import asyncio
import re
import time

import requests
from docx import Document
from docx.text.paragraph import Paragraph
import pathlib
import tqdm

import urllib3

from src.tools_for_interface_of_gigachat.Engine_of_interface import set_interface_for_job, EngineOfGigaChatInterface

urllib3.disable_warnings(urllib3.exceptions.InsecureRequestWarning)


class WordProcessor:
    def __init__(self, path_to_document: pathlib.Path):
        self.path_to_document = path_to_document
        assert path_to_document.exists(), 'Не существует данного файла'

        self.document = Document(self.path_to_document.__str__())
        self.delete_empty_paragraphs()
        self.count_of_paragraphs_to_modify = self.get_count_of_caption_text()
        assert self.count_of_paragraphs_to_modify != 0, 'Файл не может быть обработан.'

    def get_count_of_caption_text(self) -> int:
        count: int = 0

        for i, paragraph in enumerate(self.document.paragraphs):
            count += 1 if self.check_paragraph_on_suitability(paragraph, i) else 0
        return count

    @staticmethod
    def check_paragraph_on_image(paragraph: Paragraph) -> bool:
        return paragraph.runs and any(run._element.xpath('.//pic:pic') for run in paragraph.runs)

    def check_paragraph_on_suitability(self, paragraph: Paragraph, i: int) -> bool:
        if re.search('^рисунок', paragraph.text.lower().strip()) is not None:
            if i + 1 < len(self.document.paragraphs):
                if re.search('^рисунок', self.document.paragraphs[i+1].text.lower().strip()) is None:
                    return True
        return False

    def run_case(self):
        with requests.Session() as session:
            engine = set_interface_for_job(session)
            try:
                self.process_paragraphs(engine)
            finally:
                self.document.save(self.path_to_document.__str__())

    def process_paragraphs(self, engine: EngineOfGigaChatInterface):
        with tqdm.tqdm(total=self.count_of_paragraphs_to_modify) as t:
            t.set_description('Количество параграфов')
            for i, paragraph in enumerate(self.document.paragraphs):
                self.process_paragraph(engine, t, i, paragraph)

    def process_paragraph(self, engine: EngineOfGigaChatInterface, t: tqdm.tqdm, i: int, paragraph: Paragraph):
        if self.check_paragraph_on_suitability(paragraph, i):
            # использовать self.document.paragraphs[i+1].text
            time.sleep(1)
            text_from_model = engine.post_message_into_gigachat(self.document.paragraphs[i+1].text)
            self.document.paragraphs[i+1].runs[0].text = text_from_model

            t.update()

    def delete_empty_paragraphs(self):
        for paragraph in list(self.document.paragraphs):
            if paragraph.text.strip() == '' and len(paragraph.runs) == 0:
                paragraph._element.getparent().remove(paragraph._element)
        self.document.save(self.path_to_document.__str__())
